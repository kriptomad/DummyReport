"""
troubleshooter/document_extractor.py
======================================
Extracts plain text from PDF/DOCX/(legacy DOC) resolution runbooks
uploaded to the PSLD - Parts Resolution KB (troubleshooter/servicenow_
resolution_kb.py), then distills that text into a handful of "key point"
sentences — a lightweight, 100%-local stand-in for "run this through an
LLM and get back a summary."

WHY EXTRACTIVE SUMMARIZATION INSTEAD OF A REAL LLM
----------------------------------------------------
This project cannot install heavy generative-AI dependencies here (see
troubleshooter/psld_semantic_engine.py's docstring — PyTorch/
sentence-transformers fail to install because of the Windows MAX_PATH
issue caused by this repo's deeply-nested OneDrive path, and there's no
admin access to fix that). Instead, this uses TextRank-style EXTRACTIVE
summarization: split the document into sentences, score each sentence by
how similar it is to the "average" of the whole document (its TF-IDF
centrality) using the same scikit-learn machinery already used
elsewhere in this app, and keep the top N highest-scoring sentences, in
their original order. This reliably surfaces the sentences that best
represent what the document is about, without needing any external API
call or GPU.

SUPPORTED FORMATS
-----------------
- .pdf   -> pypdf (pure Python, already a project dependency after this
            change).
- .docx  -> python-docx (already a project dependency) — paragraphs and
            table cells.
- .doc   (legacy binary Word format, still used by some older KB
            runbooks) -> best-effort via Word COM automation (pywin32),
            which requires Microsoft Word to actually be installed on
            the machine running this app. If that's unavailable, a
            clear ValueError is raised asking the user to re-save the
            file as .docx or .pdf instead of silently returning nothing.
- .xlsx/.xls -> pandas/openpyxl (already project dependencies) — each
            sheet's rows are flattened into "col: value" text lines so
            TF-IDF/key-point extraction has something meaningful to
            work with.
- .txt  -> decoded directly (utf-8-sig -> utf-8 -> cp1252 -> latin-1
            fallback chain, since the real docs/ archive has files
            saved with different Windows/Brazilian-Portuguese
            encodings).
"""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import List

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:  # pragma: no cover
    PYPDF_AVAILABLE = False

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:  # pragma: no cover
    PANDAS_AVAILABLE = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:  # pragma: no cover
    SKLEARN_AVAILABLE = False

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt")

# Row cap for Excel text extraction — flattening a huge spreadsheet into
# text for TF-IDF isn't useful past a certain size and gets slow; a few
# thousand rows is already far more context than a resolution runbook
# spreadsheet realistically needs.
_MAX_EXCEL_ROWS_FOR_TEXT = 2000

_TEXT_DECODE_ENCODINGS = ("utf-8-sig", "utf-8", "cp1252", "latin-1")


class ExtractionError(ValueError):
    """Raised when a document's text can't be extracted, with a
    user-facing explanation of why (missing dependency, unsupported
    legacy format, corrupt file, etc.)."""


def _extract_pdf_text(file_bytes: bytes) -> str:
    if not PYPDF_AVAILABLE:
        raise ExtractionError(
            "PDF text extraction isn't available (the 'pypdf' package "
            "isn't installed). Run `pip install pypdf`."
        )
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()
    except Exception as e:
        raise ExtractionError(f"Could not read this PDF: {e}") from e


def _extract_docx_text(file_bytes: bytes) -> str:
    if not DOCX_AVAILABLE:
        raise ExtractionError(
            "DOCX text extraction isn't available (the 'python-docx' "
            "package isn't installed). Run `pip install python-docx`."
        )
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception as e:
        raise ExtractionError(f"Could not read this DOCX file: {e}") from e


def _extract_doc_text(file_bytes: bytes) -> str:
    """Best-effort extraction for the legacy binary .doc format via Word
    COM automation. Requires Microsoft Word installed on THIS machine
    (works when the Streamlit app runs on a Windows desktop with Office,
    which is the case here — it will NOT work on a headless Linux
    server). Raises a clear, actionable error otherwise."""
    try:
        import win32com.client  # pywin32
        import pythoncom
    except ImportError as e:
        raise ExtractionError(
            "Legacy .doc extraction needs Microsoft Word installed on "
            "this machine (via the 'pywin32' package, which isn't "
            "installed here). Easiest fix: open the file in Word and "
            "'Save As' .docx or PDF, then re-upload that instead."
        ) from e

    import tempfile
    import os as _os

    tmp_path = None
    pythoncom.CoInitialize()
    word = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(tmp_path, ReadOnly=True)
        try:
            text = doc.Content.Text
        finally:
            doc.Close(False)
        return (text or "").strip()
    except Exception as e:
        raise ExtractionError(
            f"Could not read this legacy .doc file (Word automation "
            f"failed: {e}). Try re-saving it as .docx or PDF and "
            f"re-uploading."
        ) from e
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()
        if tmp_path:
            try:
                _os.unlink(tmp_path)
            except Exception:
                pass


def _extract_excel_text(file_bytes: bytes) -> str:
    if not PANDAS_AVAILABLE:
        raise ExtractionError(
            "Excel text extraction isn't available (the 'pandas'/'openpyxl' "
            "packages aren't installed). Run `pip install pandas openpyxl`."
        )
    try:
        sheets = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    except Exception as e:
        raise ExtractionError(f"Could not read this Excel file: {e}") from e

    lines: List[str] = []
    rows_used = 0
    for sheet_name, df in sheets.items():
        if rows_used >= _MAX_EXCEL_ROWS_FOR_TEXT:
            break
        lines.append(f"[Sheet: {sheet_name}]")
        for _, row in df.head(_MAX_EXCEL_ROWS_FOR_TEXT - rows_used).iterrows():
            cells = [f"{col}: {val}" for col, val in row.items() if str(val).strip() and str(val).lower() != "nan"]
            if cells:
                lines.append(", ".join(cells))
            rows_used += 1
    return "\n".join(lines).strip()


def _extract_txt_text(file_bytes: bytes) -> str:
    for enc in _TEXT_DECODE_ENCODINGS:
        try:
            return file_bytes.decode(enc).strip()
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("utf-8", errors="replace").strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatches to the right extractor based on `filename`'s
    extension. Raises ExtractionError (a ValueError subclass) on any
    failure, with a message safe to show directly to the user."""
    ext = Path(filename or "").suffix.lower()
    if ext == ".pdf":
        return _extract_pdf_text(file_bytes)
    if ext == ".docx":
        return _extract_docx_text(file_bytes)
    if ext == ".doc":
        return _extract_doc_text(file_bytes)
    if ext in (".xlsx", ".xls"):
        return _extract_excel_text(file_bytes)
    if ext == ".txt":
        return _extract_txt_text(file_bytes)
    raise ExtractionError(
        f"Unsupported file type '{ext or '?'}'. Supported: "
        f"{', '.join(SUPPORTED_EXTENSIONS)}."
    )


_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?;:])\s+|\n+")


def _split_sentences(text: str) -> List[str]:
    candidates = [s.strip() for s in _SENTENCE_SPLIT_RE.split(text or "")]
    # Drop noise: empty lines, page-number-only lines, and sentences too
    # short to be a meaningful "key point" (e.g. bullet glyphs, headers).
    return [s for s in candidates if len(s) >= 20 and any(c.isalpha() for c in s)]


def extract_key_points(text: str, max_points: int = 8) -> List[str]:
    """
    Distills `text` (the full extracted document body) down to at most
    `max_points` sentences that best represent the document as a whole —
    a local, dependency-light "extractive summary" used in place of
    sending the document to an external LLM. Ranks sentences by their
    TF-IDF cosine similarity to the document's overall centroid (i.e.
    "how representative is this sentence of everything in the doc"),
    which tends to surface topic-sentences / conclusions / key actions
    rather than boilerplate or transitional filler. Falls back to
    "first N non-trivial sentences" if scikit-learn isn't available or
    the document is too short/sparse for TF-IDF to produce a signal.
    """
    sentences = _split_sentences(text)
    if not sentences:
        return []
    if len(sentences) <= max_points:
        return sentences

    if SKLEARN_AVAILABLE:
        try:
            vectorizer = TfidfVectorizer(max_features=4000, ngram_range=(1, 2), stop_words="english")
            matrix = vectorizer.fit_transform(sentences)
            centroid = matrix.mean(axis=0)
            # np.asarray needed: `.mean()` on a sparse matrix returns an
            # np.matrix, which cosine_similarity doesn't accept directly.
            import numpy as np
            centroid = np.asarray(centroid)
            scores = cosine_similarity(matrix, centroid).flatten()
            ranked_idx = sorted(range(len(sentences)), key=lambda i: scores[i], reverse=True)[:max_points]
            # Keep the original reading order for the picked sentences,
            # rather than dumping them sorted purely by score — reads
            # much more like a coherent summary this way.
            ranked_idx.sort()
            return [sentences[i] for i in ranked_idx]
        except ValueError:
            pass  # e.g. all-stopword text; fall through to the naive fallback

    return sentences[:max_points]
