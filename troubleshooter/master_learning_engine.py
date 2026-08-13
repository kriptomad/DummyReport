"""
🧠 MASTER LEARNING ENGINE - Sistema Unificado de Aprendizado de Troubleshooting
===============================================================================
Consolida TODOS os dados de troubleshooting e aprende padrões automaticamente
para melhorar a análise e resolução de erros.

Este módulo:
1. Processa múltiplas fontes de dados (Word, Excel, CSV)
2. Extrai padrões de erros automaticamente
3. Categoriza e prioriza problemas
4. Aprende com novos casos
5. Melhora o matching automático de soluções
"""

import pandas as pd
import os
from pathlib import Path
from docx import Document
from typing import List, Dict, Tuple, Optional
import re
from datetime import datetime
from difflib import SequenceMatcher
import json


class MasterLearningEngine:
    """Motor de aprendizado mestre para troubleshooting"""

    def __init__(self):
        # Resolve paths relative to the repo, not to a specific developer's
        # machine, so this works on any environment.
        self.base_path = str(Path(__file__).resolve().parent.parent)
        self.downloads_path = str(Path.home() / "Downloads")
        self.assets_path = os.path.join(self.base_path, 'assets')

        # Base de conhecimento em memória
        self.knowledge_base = {
            'errors': [],
            'patterns': {},
            'categories': set(),
            'keywords': {},
            'solutions': {},
            'stats': {}
        }

        # Configuração de colunas padrão
        self.standard_columns = [
            'Categoria',
            'Mensagem de erro / padrão identificado',
            'Significado provável',
            'Precisa usar a Tariff Pool Query?',
            'Como validar',
            'Ação recomendada',
            'Responsável sugerido'
        ]

    def learn_from_docx(self, filepath: str) -> List[Dict]:
        """Aprende com documentos Word"""

        print(f"\n📄 Aprendendo de Word: {os.path.basename(filepath)}")

        if not os.path.exists(filepath):
            print(f"   ⚠️ Arquivo não encontrado")
            return []

        try:
            doc = Document(filepath)
            learned_data = []

            # Extrair tabelas
            for idx, table in enumerate(doc.tables):
                headers, data = None, []

                for i, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]

                    if i == 0:
                        headers = row_data
                    elif headers and len(row_data) == len(headers):
                        data.append(dict(zip(headers, row_data)))

                if data:
                    learned_data.extend(data)

            # Extrair texto para análise de padrões
            full_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            patterns = self._extract_error_patterns(full_text)

            print(f"   ✅ {len(learned_data)} registros + {len(patterns)} padrões extraídos")

            return learned_data + patterns

        except Exception as e:
            print(f"   ❌ Erro: {str(e)}")
            return []

    def learn_from_excel(self, filepath: str) -> List[Dict]:
        """Aprende com arquivos Excel"""

        print(f"\n📊 Aprendendo de Excel: {os.path.basename(filepath)}")

        if not os.path.exists(filepath):
            print(f"   ⚠️ Arquivo não encontrado")
            return []

        try:
            excel = pd.ExcelFile(filepath)
            all_data = []

            for sheet in excel.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet)
                df = df.dropna(how='all').fillna('')

                # Converter para dicionários
                records = df.to_dict('records')
                all_data.extend(records)

                print(f"   ✅ {sheet}: {len(records)} registros")

            return all_data

        except Exception as e:
            print(f"   ❌ Erro: {str(e)}")
            return []

    def _extract_error_patterns(self, text_lines: List[str]) -> List[Dict]:
        """Extrai padrões de erro do texto"""

        patterns = []
        error_indicators = [
            r'error[:|\s]+(.*)',
            r'failed[:|\s]+(.*)',
            r'invalid[:|\s]+(.*)',
            r'missing[:|\s]+(.*)',
            r'not found[:|\s]+(.*)',
            r'unable to[:|\s]+(.*)',
            r'cannot[:|\s]+(.*)',
            r'duplicate[:|\s]+(.*)',
            r'expired[:|\s]+(.*)',
        ]

        for line in text_lines:
            line_lower = line.lower()

            for pattern in error_indicators:
                match = re.search(pattern, line_lower, re.IGNORECASE)
                if match:
                    patterns.append({
                        'Error Message': line,
                        'Category': 'Auto-Detected Pattern',
                        'Description': match.group(1) if match.groups() else line,
                        'Source': 'Pattern Extraction'
                    })
                    break

        return patterns

    def standardize_data(self, raw_data: List[Dict]) -> pd.DataFrame:
        """Padroniza dados de múltiplas fontes"""

        # Mapeamento de colunas
        column_map = {
            'Error Message': 'Mensagem de erro / padrão identificado',
            'Error Pattern': 'Mensagem de erro / padrão identificado',
            'Error': 'Mensagem de erro / padrão identificado',
            'Message': 'Mensagem de erro / padrão identificado',

            'Meaning': 'Significado provável',
            'Description': 'Significado provável',
            'Root Cause': 'Significado provável',
            'Cause': 'Significado provável',

            'Validation': 'Como validar',
            'Validation Required': 'Como validar',
            'How to Validate': 'Como validar',

            'Action': 'Ação recomendada',
            'Action Required': 'Ação recomendada',
            'Recommended Action': 'Ação recomendada',
            'Resolution': 'Ação recomendada',

            'Owner': 'Responsável sugerido',
            'Responsible': 'Responsável sugerido',
            'Team': 'Responsável sugerido',

            'Category': 'Categoria',
            'Error Category': 'Categoria',
            'Type': 'Categoria',
            'Error Type': 'Categoria',

            'Tariff Query': 'Precisa usar a Tariff Pool Query?',
            'Needs Tariff Query': 'Precisa usar a Tariff Pool Query?',
        }

        # Converter para DataFrame
        df = pd.DataFrame(raw_data)

        # Renomear colunas
        df.rename(columns=column_map, inplace=True)

        # Garantir colunas padrão
        for col in self.standard_columns:
            if col not in df.columns:
                df[col] = ''

        # Limpar dados
        df = df.fillna('')

        return df

    def analyze_and_learn(self, df: pd.DataFrame) -> Dict:
        """Analisa dados e aprende padrões"""

        print("\n🧠 Analisando e aprendendo padrões...")

        error_col = 'Mensagem de erro / padrão identificado'

        # Análise de palavras-chave
        keywords = ['missing', 'invalid', 'not found', 'failed', 'error',
                   'unable', 'cannot', 'denied', 'expired', 'duplicate',
                   'tariff', 'rate', 'lane', 'origin', 'destination',
                   'equipment', 'service', 'charge', 'carrier']

        all_errors = ' '.join(df[error_col].astype(str).values).lower()

        keyword_freq = {}
        for kw in keywords:
            count = all_errors.count(kw)
            if count > 0:
                keyword_freq[kw] = count

        # Análise de categorias
        category_dist = df['Categoria'].value_counts().to_dict() if 'Categoria' in df.columns else {}

        # Padrões de erro por tipo
        error_types = {}
        for _, row in df.iterrows():
            msg = str(row[error_col]).lower()

            # Classificar tipo
            if 'missing' in msg or 'not found' in msg:
                error_type = 'Missing Data'
            elif 'invalid' in msg:
                error_type = 'Invalid Data'
            elif 'duplicate' in msg:
                error_type = 'Duplicate'
            elif 'expired' in msg or 'date' in msg:
                error_type = 'Date/Expiration'
            elif 'tariff' in msg or 'rate' in msg:
                error_type = 'Tariff/Rate'
            elif 'lane' in msg or 'route' in msg:
                error_type = 'Routing/Lane'
            else:
                error_type = 'Other'

            error_types[error_type] = error_types.get(error_type, 0) + 1

        # Análise de soluções mais comuns
        solutions = df['Ação recomendada'].value_counts().head(10).to_dict() if 'Ação recomendada' in df.columns else {}

        # Responsáveis mais frequentes
        owners = df['Responsável sugerido'].value_counts().head(10).to_dict() if 'Responsável sugerido' in df.columns else {}

        tariff_col = 'Precisa usar a Tariff Pool Query?'
        tariff_query_needed = (
            len(df[df[tariff_col].astype(str).str.upper() == 'SIM'])
            if tariff_col in df.columns else 0
        )

        analysis = {
            'total_errors': len(df),
            'unique_patterns': df[error_col].nunique(),
            'keywords': keyword_freq,
            'categories': category_dist,
            'error_types': error_types,
            'top_solutions': solutions,
            'top_owners': owners,
            'tariff_query_needed': tariff_query_needed
        }

        # Atualizar base de conhecimento
        self.knowledge_base['patterns'] = error_types
        self.knowledge_base['keywords'] = keyword_freq
        self.knowledge_base['categories'] = set(category_dist.keys())
        self.knowledge_base['stats'] = analysis

        print(f"   ✅ Aprendizado concluído: {analysis['total_errors']} erros analisados")

        return analysis

    def smart_match(self, error_message: str, threshold: float = 0.6) -> List[Dict]:
        """
        Matching inteligente de erros usando análise de similaridade

        Args:
            error_message: Mensagem de erro a ser analisada
            threshold: Limite de similaridade (0.0 a 1.0)

        Returns:
            Lista de soluções ordenadas por relevância
        """

        # Carregar base de dados consolidada
        db_path = os.path.join(self.assets_path, 'dummytroubleshoot_complete.xlsx')

        if not os.path.exists(db_path):
            return []

        df = pd.read_excel(db_path)

        matches = []
        error_lower = error_message.lower()

        for idx, row in df.iterrows():
            pattern = str(row.get('Mensagem de erro / padrão identificado', '')).lower()

            # Calcular similaridade
            similarity = SequenceMatcher(None, error_lower, pattern).ratio()

            # Verificar palavras-chave comuns
            keyword_boost = 0
            for keyword in self.knowledge_base.get('keywords', {}).keys():
                if keyword in error_lower and keyword in pattern:
                    keyword_boost += 0.1

            final_score = min(similarity + keyword_boost, 1.0)

            if final_score >= threshold:
                matches.append({
                    'score': final_score,
                    'pattern': row.get('Mensagem de erro / padrão identificado', ''),
                    'category': row.get('Categoria', ''),
                    'meaning': row.get('Significado provável', ''),
                    'validation': row.get('Como validar', ''),
                    'action': row.get('Ação recomendada', ''),
                    'owner': row.get('Responsável sugerido', ''),
                    'needs_tariff': row.get('Precisa usar a Tariff Pool Query?', '')
                })

        # Ordenar por score
        matches.sort(key=lambda x: x['score'], reverse=True)

        return matches[:5]  # Top 5 matches

    def save_learned_knowledge(self, output_filename: str = 'master_troubleshooting_db.xlsx'):
        """Salva base de conhecimento consolidada"""

        output_path = os.path.join(self.assets_path, output_filename)

        # Carregar todos os arquivos existentes
        all_files = [
            'dummytroubleshoot_complete.xlsx',
            'master_data_troubleshooting_consolidated.xlsx',
            'stepsdummy.xlsx'
        ]

        all_data = []

        for filename in all_files:
            filepath = os.path.join(self.assets_path, filename)
            if os.path.exists(filepath):
                try:
                    df = pd.read_excel(filepath)
                    df['Source_File'] = filename
                    all_data.append(df)
                    print(f"   ✅ Carregado: {filename}")
                except Exception as e:
                    print(f"   ⚠️ Erro em {filename}: {str(e)}")

        if not all_data:
            print("   ❌ Nenhum dado para consolidar")
            return None

        # Mesclar tudo
        df_master = pd.concat(all_data, ignore_index=True)

        # Padronizar
        df_master = self.standardize_data(df_master.to_dict('records'))

        # Remover duplicatas
        before = len(df_master)
        df_master = df_master.drop_duplicates(
            subset=['Mensagem de erro / padrão identificado'],
            keep='first'
        )
        after = len(df_master)

        print(f"\n💾 Salvando base mestre de conhecimento...")
        print(f"   Total: {after} registros únicos")
        print(f"   Duplicatas removidas: {before - after}")

        # Salvar com múltiplas abas
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Aba principal
            df_master.to_excel(writer, sheet_name='All Errors', index=False)

            # Por categoria
            for cat in df_master['Categoria'].unique():
                if cat and cat != '':
                    df_cat = df_master[df_master['Categoria'] == cat]
                    sheet_name = str(cat).replace('/', '-').replace('\\', '-')[:31]
                    df_cat.to_excel(writer, sheet_name=sheet_name, index=False)

            # Estatísticas
            analysis = self.analyze_and_learn(df_master)

            stats_df = pd.DataFrame({
                'Métrica': [
                    'Total de Erros',
                    'Padrões Únicos',
                    'Categorias',
                    'Precisam Tariff Query'
                ],
                'Valor': [
                    analysis['total_errors'],
                    analysis['unique_patterns'],
                    len(analysis['categories']),
                    analysis['tariff_query_needed']
                ]
            })
            stats_df.to_excel(writer, sheet_name='Statistics', index=False)

            # Keywords
            if analysis['keywords']:
                kw_df = pd.DataFrame([
                    {'Keyword': k, 'Frequência': v}
                    for k, v in sorted(analysis['keywords'].items(),
                                      key=lambda x: x[1], reverse=True)
                ])
                kw_df.to_excel(writer, sheet_name='Keywords', index=False)

            # Tipos de erro
            if analysis['error_types']:
                types_df = pd.DataFrame([
                    {'Tipo de Erro': k, 'Quantidade': v}
                    for k, v in sorted(analysis['error_types'].items(),
                                      key=lambda x: x[1], reverse=True)
                ])
                types_df.to_excel(writer, sheet_name='Error Types', index=False)

        print(f"   ✅ Salvo: {output_path}")

        # Salvar metadata como JSON
        metadata = {
            'last_updated': datetime.now().isoformat(),
            'total_errors': analysis['total_errors'],
            'unique_patterns': analysis['unique_patterns'],
            'categories': list(analysis['categories']),
            'keywords': analysis['keywords'],
            'error_types': analysis['error_types']
        }

        metadata_path = os.path.join(self.assets_path, 'knowledge_base_metadata.json')
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2, ensure_ascii=False)

        print(f"   ✅ Metadata salvo: {metadata_path}")

        return df_master

    def process_new_files(self, file_list: List[str]) -> pd.DataFrame:
        """Processa novos arquivos de troubleshooting"""

        print("=" * 80)
        print("🚀 INICIANDO APRENDIZADO DE NOVOS ARQUIVOS")
        print("=" * 80)

        all_learned_data = []

        for filepath in file_list:
            if not os.path.exists(filepath):
                print(f"\n⚠️ Arquivo não encontrado: {filepath}")
                continue

            ext = os.path.splitext(filepath)[1].lower()

            if ext == '.docx':
                data = self.learn_from_docx(filepath)
            elif ext in ['.xlsx', '.xls']:
                data = self.learn_from_excel(filepath)
            else:
                print(f"\n⚠️ Formato não suportado: {ext}")
                continue

            all_learned_data.extend(data)

        if not all_learned_data:
            print("\n❌ Nenhum dado aprendido!")
            return None

        # Padronizar
        df = self.standardize_data(all_learned_data)

        # Analisar
        analysis = self.analyze_and_learn(df)

        # Gerar relatório
        self._print_learning_report(analysis)

        return df

    def _print_learning_report(self, analysis: Dict):
        """Imprime relatório de aprendizado"""

        print("\n" + "=" * 80)
        print("📚 RELATÓRIO DE APRENDIZADO")
        print("=" * 80)

        print(f"\n📊 RESUMO:")
        print(f"   • Total de erros: {analysis['total_errors']}")
        print(f"   • Padrões únicos: {analysis['unique_patterns']}")
        print(f"   • Taxa de duplicação: {((1 - analysis['unique_patterns']/analysis['total_errors']) * 100):.1f}%")
        print(f"   • Erros que precisam Tariff Query: {analysis['tariff_query_needed']}")

        print(f"\n📂 CATEGORIAS ({len(analysis['categories'])}):")
        for cat, count in sorted(analysis['categories'].items(),
                                key=lambda x: x[1], reverse=True)[:10]:
            pct = (count / analysis['total_errors']) * 100
            print(f"   • {cat}: {count} ({pct:.1f}%)")

        print(f"\n🔍 KEYWORDS IMPORTANTES:")
        for kw, freq in sorted(analysis['keywords'].items(),
                              key=lambda x: x[1], reverse=True)[:15]:
            print(f"   • '{kw}': {freq}x")

        print(f"\n🏷️ TIPOS DE ERRO:")
        for error_type, count in sorted(analysis['error_types'].items(),
                                       key=lambda x: x[1], reverse=True):
            pct = (count / analysis['total_errors']) * 100
            print(f"   • {error_type}: {count} ({pct:.1f}%)")

        print("\n" + "=" * 80)


def main():
    """Função principal para executar o aprendizado"""

    engine = MasterLearningEngine()

    # Arquivos para processar (coloque os arquivos na pasta Downloads do usuário)
    new_files = [
        os.path.join(engine.downloads_path, 'Automated TM Master Data Errors.docx'),
        os.path.join(engine.downloads_path, 'data master troubleshooting.xlsx'),
    ]

    # Processar novos arquivos
    df_new = engine.process_new_files(new_files)

    if df_new is not None:
        # Salvar temporariamente
        temp_output = os.path.join(engine.assets_path, 'temp_new_learning.xlsx')
        df_new.to_excel(temp_output, index=False)
        print(f"\n💾 Dados novos salvos: {temp_output}")

    # Consolidar TUDO
    df_master = engine.save_learned_knowledge()

    if df_master is not None:
        print("\n" + "=" * 80)
        print("✅ APRENDIZADO CONCLUÍDO COM SUCESSO!")
        print("=" * 80)
        print(f"\n📚 Base de conhecimento consolidada:")
        print(f"   • {len(df_master)} erros catalogados")
        print(f"   • {df_master['Categoria'].nunique()} categorias")
        print(f"   • {df_master['Responsável sugerido'].nunique()} responsáveis")

        # Teste de matching
        print("\n🧪 TESTE DE MATCHING INTELIGENTE:")
        test_error = "Missing tariff rate for origin destination"
        matches = engine.smart_match(test_error)

        if matches:
            print(f"\n   Erro teste: '{test_error}'")
            print(f"   Encontrados {len(matches)} matches:")
            for i, match in enumerate(matches, 1):
                print(f"\n   {i}. Score: {match['score']:.2%}")
                print(f"      Padrão: {match['pattern'][:60]}...")
                print(f"      Ação: {match['action'][:60]}...")

        print("\n" + "=" * 80)


if __name__ == "__main__":
    main()

